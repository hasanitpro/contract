from __future__ import annotations

from io import BytesIO
from typing import Dict
from docx import Document
from docx.shared import Cm
from docx.oxml import OxmlElement

# ============================================================
# Core public API
# ============================================================

def generate_docx_from_template(
    template_path: str,
    ctx: Dict[str, str],
) -> bytes:
    """
    Generate final DOCX by replacing placeholders in a lawyer-approved template.

    IMPORTANT CONCEPT:
    ------------------
    - This function ONLY renders a DOCX.
    - It does NOT decide legal logic.
    - All legal/business decisions must already be resolved
      and passed in via `ctx`.

    Placeholder rules:
    ------------------
    - Inline placeholders like [VERMIETER_NAME] can appear anywhere
    - Block placeholders like [SR_BLOCK] MUST be alone in a paragraph
    - Empty block values REMOVE the paragraph completely
    """

    # Load the DOCX template
    doc = Document(template_path)

    # 1️⃣ Replace inline placeholders in normal paragraphs
    for paragraph in doc.paragraphs:
        replace_inline_placeholders(paragraph, ctx)

    # 2️⃣ Replace inline placeholders inside tables
    replace_placeholders_in_tables(doc, ctx)

    # 3️⃣ Replace or remove block placeholders
    replace_block_placeholders(doc, ctx)

    # Save final document to memory (no disk write here)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ============================================================
# Context builder (VERY IMPORTANT)
# ============================================================

def build_placeholder_context(maskA: dict, maskB: dict) -> Dict[str, str]:
    """
    Build ONE flat placeholder dictionary from Mask A + Mask B.

    WHY THIS EXISTS:
    ----------------
    - Keeps legal logic separate from DOCX rendering
    - Makes debugging easy
    - Allows lawyers to reason about contract output

    This function is where:
    - Mask A (facts) are mapped
    - Mask B (legal decisions) are resolved
    """

    ctx: Dict[str, str] = {}

    # ----------------------------
    # Mask A – factual data
    # ----------------------------
    ctx["VERMIETER_NAME"] = maskA.get("vermieter_name", "")
    ctx["MIETER_NAME"] = maskA.get("mieter_name", "")
    ctx["OBJEKT_ADRESSE"] = maskA.get("objekt_adresse", "")

    # ----------------------------
    # Mask B – simple fields
    # ----------------------------
    ctx["RENT_START_DATE"] = maskB.get("mietbeginn", "")
    
    # ----------------------------
    # Mask B – Kündigung / Laufzeit logic (STEP 6)
    # ----------------------------
    ctx.update(resolve_kuendigung(maskB))
    # 👆 THIS is the correct place for:
    # placeholders.update(resolve_kuendigung(maskB))

    return ctx


# ============================================================
# Inline placeholder replacement
# ============================================================

def replace_inline_placeholders(paragraph, ctx: Dict[str, str]) -> None:
    """
    Replace placeholders inside a paragraph.

    WHY THIS IS COMPLEX:
    -------------------
    Word splits text into multiple "runs", which can break
    simple string replacement. This function safely rebuilds
    the paragraph if needed.
    """

    if not paragraph.text:
        return

    full_text = paragraph.text
    replaced = False

    for key, value in ctx.items():
        # Inline replacement ONLY supports strings
        if not isinstance(value, str):
            continue

        placeholder = f"[{key}]"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)
            replaced = True


    if not replaced:
        return

    # Clear all runs and rebuild paragraph safely
    clear_paragraph(paragraph)
    paragraph.add_run(full_text)


# ============================================================
# Table handling
# ============================================================

def replace_placeholders_in_tables(doc, ctx: Dict[str, str]) -> None:
    """
    Apply inline placeholder replacement inside tables.
    Word tables behave differently than normal paragraphs.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_inline_placeholders(paragraph, ctx)

# ============================================================
# Special table insertion: § 5 Miete / Betriebskosten
# ============================================================

def insert_miete_bk_table_after_paragraph(doc, paragraph, rows):
    """
    Insert a formatted 2-column table (Beschreibung | Betrag)
    directly after the placeholder paragraph.
    """

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.autofit = False

    # Header
    table.cell(0, 0).text = "Beschreibung"
    table.cell(0, 1).text = "Betrag (EUR)"

    # Data rows
    for i, (label, amount) in enumerate(rows, start=1):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = amount

    # Column widths (legal-doc friendly)
    set_table_column_widths(table, [12.5, 4.0])

    # Borders
    set_table_borders(table)

    # ✅ Vertical alignment (CENTER for all cells)
    for row in table.rows:
        for cell in row.cells:
            set_cell_vertical_alignment(cell, "center")

    # Row heights
    set_row_height(table.rows[0], 0.95)   # header
    for row in table.rows[1:-1]:
        set_row_height(row, 1.25)          # normal rows
        set_row_height(table.rows[-1], 0.95) # total row

    # Insert table after placeholder
    paragraph._p.addnext(table._tbl)



# ============================================================
# Block placeholder handling
# ============================================================

def replace_block_placeholders(doc, ctx: Dict[str, str]) -> None:
    """
    Replace placeholders that MUST be alone in a paragraph.
    """

    block_keys = [
        "MPB_BLOCK",
        "WEG_BLOCK",
        "SR_BLOCK",
        "ANNEX_LIST",
        "CLAUSE_KUENDIGUNGSAUSSCHLUSS",
        "CLAUSE_NEBENKOSTEN",
        "STAFFEL_BLOCK",
        "ANNEX_BLOCK",
        "CLAUSE_UNTERVERMIETUNG",
        "CLAUSE_TIERHALTUNG",
        "CLAUSE_SCHOENHEITSREPARATUREN",
        "CLAUSE_ENDRUECKGABE",
        "CLAUSE_DATENVERARBEITUNG_ENERGIE_ANLAGEN",
    ]

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()

        # ====================================================
        # SPECIAL CASE: Real DOCX table for § 5
        # ====================================================
        if text == "[MIETE_BK_TABELLE]":
            rows = ctx.get("MIETE_BK_TABELLE", [])

            if rows:
                insert_miete_bk_table_after_paragraph(doc, paragraph, rows)

            delete_paragraph(paragraph)
            continue

        # ====================================================
        # Normal block placeholders
        # ====================================================
        for key in block_keys:
            placeholder = f"[{key}]"
            if text == placeholder:
                value = ctx.get(key, "")

                if isinstance(value, str) and value.strip():
                    clear_paragraph(paragraph)
                    paragraph.add_run(value)
                else:
                    delete_paragraph(paragraph)

# ============================================================
# Low-level helpers (DO NOT TOUCH)
# ============================================================

def clear_paragraph(paragraph) -> None:
    """
    Remove all runs from a paragraph.
    This is safer than deleting the paragraph.
    """
    for run in paragraph.runs:
        run.text = ""


def delete_paragraph(paragraph) -> None:
    """
    Completely remove a paragraph from the document.

    WARNING:
    --------
    This uses Word XML internals.
    Only use this exact implementation.
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

from docx.oxml import OxmlElement, ns

def set_table_borders(table):
    """
    Apply full borders to a table (Word-compatible, safe).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement("w:tblBorders")

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(ns.qn("w:val"), "single")
        element.set(ns.qn("w:sz"), "8")       # border thickness
        element.set(ns.qn("w:space"), "0")
        element.set(ns.qn("w:color"), "000000")
        borders.append(element)

    tblPr.append(borders)

def set_table_column_widths(table, widths_cm):
    """
    widths_cm = list of column widths in centimeters
    """
    table.autofit = False

    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)

from docx.oxml import OxmlElement, ns

def set_cell_vertical_alignment(cell, align="center"):
    """
    Vertically align content inside a table cell.
    align: 'top' | 'center' | 'bottom'
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    vAlign = OxmlElement("w:vAlign")
    vAlign.set(ns.qn("w:val"), align)

    tcPr.append(vAlign)



def set_row_height(row, height_cm: float):
    """
    Set exact row height in centimeters.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()

    trHeight = OxmlElement("w:trHeight")
    trHeight.set(ns.qn("w:val"), str(int(height_cm * 567)))  # cm → twips
    trHeight.set(ns.qn("w:hRule"), "exact")

    trPr.append(trHeight)

# ============================================================
# Legal logic: Kündigung & Laufzeit (STEP 6)
# ============================================================

def resolve_kuendigung(mask_b: dict) -> Dict[str, str]:
    """
    Resolve contract duration & termination rules.

    LEGAL PRINCIPLES:
    -----------------
    - Unbefristet → Kündigung mit Fristen
    - Befristet → Kündigung vor Ablauf ausgeschlossen (§ 575 BGB)
    """

    result: Dict[str, str] = {}

    vertragsart = mask_b.get("vertragsart")

    if vertragsart == "unbefristet":
        result["CONTRACT_TYPE"] = "unbefristet"
        result["TENANCY_END_DATE"] = ""
        result["TERMINATION_NOTICE_Mieter"] = (
            f"{mask_b.get('kuendigungsfrist_mieter_monate', 3)} Monate"
        )
        result["TERMINATION_NOTICE_Vermieter"] = (
            f"{mask_b.get('kuendigungsfrist_vermieter_monate', 3)} Monate"
        )
        result["BEFRISTUNG_REASON"] = ""

    elif vertragsart == "befristet":
        result["CONTRACT_TYPE"] = "befristet"
        result["TENANCY_END_DATE"] = mask_b.get("befristet_enddatum", "")
        result["TERMINATION_NOTICE_Mieter"] = "Vor Ablauf ausgeschlossen"
        result["TERMINATION_NOTICE_Vermieter"] = "Vor Ablauf ausgeschlossen"
        result["BEFRISTUNG_REASON"] = (
            "Befristungsgrund: "
            + mask_b.get("befristung_grund", "")
        )

    else:
        raise ValueError(
            "vertragsart muss 'befristet' oder 'unbefristet' sein"
        )

    return result
